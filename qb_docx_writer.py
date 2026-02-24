"""
Question Bank Word Document Generator
Produces a professional .docx with Part A (SA+MCQ) and Part B, unit-wise
"""
import io
import base64
import re
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

# ─── Color palette ────────────────────────────────────────
C_DARK_BLUE   = RGBColor(0x1A, 0x23, 0x7E)
C_MID_BLUE    = RGBColor(0x28, 0x35, 0x93)
C_LIGHT_BLUE  = RGBColor(0xE8, 0xEA, 0xF6)
C_TEAL        = RGBColor(0x00, 0x6D, 0x77)
C_GOLD        = RGBColor(0xE9, 0xC4, 0x6A)
C_WHITE       = RGBColor(0xFF, 0xFF, 0xFF)
C_LIGHT_GREY  = RGBColor(0xF5, 0xF5, 0xF5)
C_BORDER      = RGBColor(0xC5, 0xCA, 0xE9)

UNIT_NAMES = {
    '1': 'Unit 1 – Sequential Circuit Design',
    '2': 'Unit 2 – Asynchronous Sequential Circuit Design',
    '3': 'Unit 3 – Testing',
    '4': 'Unit 4 – PLD Design',
    '5': 'Unit 5 – Digital Systems & Programming Tools',
}
SOURCE_LABELS = {'ct': 'Cycle Test', 'es': 'End Semester', 'aqa': 'AQA / Question Bank'}
TYPE_LABELS   = {'SA': 'Part A – Short Answer Questions (2 Marks)',
                 'MCQ': 'Part A – MCQ Questions (2 Marks)',
                 'PARTB': 'Part B Questions (8 / 16 Marks)'}

# ─── XML helpers ─────────────────────────────────────────
def set_cell_bg(cell, hex_color):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def set_cell_borders(cell, color='C5CAE9', size=4):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side in ('top','left','bottom','right'):
        border = OxmlElement(f'w:{side}')
        border.set(qn('w:val'),   'single')
        border.set(qn('w:sz'),    str(size))
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), color)
        tcBorders.append(border)
    tcPr.append(tcBorders)

def set_table_borders(table, color='C5CAE9'):
    for row in table.rows:
        for cell in row.cells:
            set_cell_borders(cell, color)

def set_cell_margins(cell, top=60, bottom=60, left=100, right=100):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for side, val in [('top',top),('bottom',bottom),('left',left),('right',right)]:
        m = OxmlElement(f'w:{side}')
        m.set(qn('w:w'), str(val))
        m.set(qn('w:type'), 'dxa')
        tcMar.append(m)
    tcPr.append(tcMar)

def add_paragraph_border_bottom(para, color='1A237E', size=12):
    pPr  = para._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bdr  = OxmlElement('w:bottom')
    bdr.set(qn('w:val'),   'single')
    bdr.set(qn('w:sz'),    str(size))
    bdr.set(qn('w:space'), '1')
    bdr.set(qn('w:color'), color)
    pBdr.append(bdr)
    pPr.append(pBdr)

def set_col_width(table, col_idx, width_dxa):
    for row in table.rows:
        if col_idx < len(row.cells):
            tc = row.cells[col_idx]._tc
            tcPr = tc.get_or_add_tcPr()
            tcW  = tcPr.find(qn('w:tcW'))
            if tcW is None:
                tcW = OxmlElement('w:tcW')
                tcPr.append(tcW)
            tcW.set(qn('w:w'),    str(width_dxa))
            tcW.set(qn('w:type'), 'dxa')

# ─── Document setup ──────────────────────────────────────
def create_document():
    doc = Document()
    # A4 page
    for section in doc.sections:
        section.page_width  = Cm(21)
        section.page_height = Cm(29.7)
        section.left_margin = section.right_margin = Cm(2)
        section.top_margin  = section.bottom_margin = Cm(1.8)

    # Default styles
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(11)
    style.paragraph_format.space_after = Pt(0)
    return doc

# ─── Header block ────────────────────────────────────────
def add_document_header(doc, subject, code, dept, semester, generated_from):
    # Title bar
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f'{code} – {subject}')
    run.bold      = True
    run.font.size = Pt(16)
    run.font.color.rgb = C_DARK_BLUE
    add_paragraph_border_bottom(p, '1A237E', 12)

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run(f'QUESTION BANK | Department of {dept or "ECE"} | {semester or "Odd Semester"}')
    r2.font.size  = Pt(10)
    r2.font.color.rgb = C_TEAL
    r2.bold = True

    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r3 = p3.add_run(f'Source Files: {generated_from}')
    r3.font.size  = Pt(9)
    r3.font.color.rgb = RGBColor(0x78, 0x78, 0x78)

    doc.add_paragraph()  # spacer

# ─── CO statements table ─────────────────────────────────
def add_co_table(doc, co_statements):
    if not co_statements: return
    p = doc.add_paragraph()
    r = p.add_run('Course Outcomes (CO)')
    r.bold = True; r.font.size = Pt(11); r.font.color.rgb = C_DARK_BLUE

    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    for c, t in zip(hdr, ['CO', 'Statement']):
        c.text = t
        c.paragraphs[0].runs[0].bold = True
        c.paragraphs[0].runs[0].font.color.rgb = C_WHITE
        set_cell_bg(c, '1A237E')
        set_cell_margins(c)

    for co, stmt in co_statements.items():
        row = table.add_row()
        row.cells[0].text = co
        row.cells[1].text = stmt
        row.cells[0].paragraphs[0].runs[0].bold = True
        set_cell_bg(row.cells[0], 'E8EAF6')
        set_cell_margins(row.cells[0])
        set_cell_margins(row.cells[1])

    set_col_width(table, 0, 800)
    set_col_width(table, 1, 7500)
    set_table_borders(table)
    doc.add_paragraph()

# ─── Section heading ──────────────────────────────────────
def add_section_heading(doc, title, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(4)
    if level == 1:
        r = p.add_run(title.upper())
        r.bold = True; r.font.size = Pt(13); r.font.color.rgb = C_DARK_BLUE
        add_paragraph_border_bottom(p, '1A237E', 8)
    elif level == 2:
        r = p.add_run(title)
        r.bold = True; r.font.size = Pt(11); r.font.color.rgb = C_MID_BLUE
        add_paragraph_border_bottom(p, '283593', 4)
    else:
        r = p.add_run(f'▌ {title}')
        r.bold = True; r.font.size = Pt(10.5); r.font.color.rgb = C_TEAL

# ─── Image inserter ──────────────────────────────────────
def insert_image_from_datauri(doc_or_cell, data_uri, max_width=Inches(3.5)):
    """Insert base64 image into a paragraph"""
    try:
        # Extract base64 data
        match = re.match(r'data:(image/\w+);base64,(.+)', data_uri, re.S)
        if not match: return None
        mime, b64data = match.groups()
        img_bytes = base64.b64decode(b64data)
        stream = io.BytesIO(img_bytes)

        if hasattr(doc_or_cell, 'add_paragraph'):
            # It's a document
            p = doc_or_cell.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run()
            run.add_picture(stream, width=max_width)
            return p
        else:
            # It's a cell – add to its last paragraph
            para = doc_or_cell.paragraphs[-1]
            run  = para.add_run()
            run.add_picture(stream, width=max_width)
            return para
    except Exception as e:
        print(f'[IMG WARN] {e}')
        return None

# ─── Question row builder ─────────────────────────────────
def add_question_row(table, q, row_num, show_options=True):
    """Add one question as a table row"""
    row = table.add_row()
    cells = row.cells

    # Alternate row shading
    if row_num % 2 == 0:
        for c in cells: set_cell_bg(c, 'F5F7FF')

    for c in cells: set_cell_margins(c, top=80, bottom=80, left=100, right=100)

    # Col 0: Q.No
    p0 = cells[0].paragraphs[0]
    r0 = p0.add_run(str(q.qno or row_num))
    r0.font.size = Pt(10); r0.bold = True
    p0.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Col 1: CO
    p1 = cells[1].paragraphs[0]
    r1 = p1.add_run(q.co)
    r1.font.size = Pt(10); r1.bold = True; r1.font.color.rgb = C_DARK_BLUE
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Col 2: BTL (K-level)
    p2 = cells[2].paragraphs[0]
    r2 = p2.add_run(q.btl)
    r2.font.size = Pt(10); r2.bold = True
    # Color by BTL
    btl_colors = {'K1': RGBColor(0x2E,0x7D,0x32), 'K2': RGBColor(0x15,0x65,0xC0),
                  'K3': RGBColor(0xE6,0x51,0x00), 'K4': RGBColor(0x6A,0x1B,0x9A),
                  'K5': RGBColor(0xC6,0x28,0x28), 'K6': RGBColor(0x37,0x47,0x4F)}
    r2.font.color.rgb = btl_colors.get(q.btl, C_DARK_BLUE)
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Col 3: Question text (with images)
    q_cell = cells[3]
    q_cell.paragraphs[0].clear()
    p_q = q_cell.paragraphs[0]

    # Main question text
    lines = q.text.split('\n')
    for i, line in enumerate(lines):
        if not line.strip(): continue
        if i == 0:
            run = p_q.add_run(line.strip())
            run.font.size = Pt(10.5)
        else:
            p_extra = q_cell.add_paragraph()
            p_extra.paragraph_format.space_before = Pt(2)
            run = p_extra.add_run(line.strip())
            run.font.size = Pt(10.5)

    # MCQ options
    if show_options and q.type == 'MCQ' and q.options:
        opt_labels = ['a)', 'b)', 'c)', 'd)']
        for idx, opt in enumerate(q.options[:4]):
            if not opt: continue
            p_opt = q_cell.add_paragraph()
            p_opt.paragraph_format.left_indent = Pt(14)
            p_opt.paragraph_format.space_before = Pt(1)
            r_lbl = p_opt.add_run(f'{opt_labels[idx]} ')
            r_lbl.bold = True; r_lbl.font.size = Pt(10)
            r_lbl.font.color.rgb = C_TEAL
            r_txt = p_opt.add_run(opt.strip())
            r_txt.font.size = Pt(10)

    # Embedded images
    for img_uri in q.images:
        try:
            match = re.match(r'data:(image/\w+);base64,(.+)', img_uri, re.S)
            if not match: continue
            mime, b64data = match.groups()
            img_bytes = base64.b64decode(b64data)
            stream = io.BytesIO(img_bytes)
            p_img = q_cell.add_paragraph()
            p_img.paragraph_format.space_before = Pt(4)
            run_img = p_img.add_run()
            run_img.add_picture(stream, width=Inches(2.8))
        except Exception as e:
            p_err = q_cell.add_paragraph()
            p_err.add_run(f'[Image: could not embed – {e}]').font.size = Pt(8)

    # Col 4: Marks
    p4 = cells[4].paragraphs[0]
    r4 = p4.add_run(str(q.marks))
    r4.font.size = Pt(10.5); r4.bold = True
    p4.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Col 5: Source
    src_label = {'ct':'CT', 'es':'ES', 'aqa':'AQA'}.get(q.source, q.source.upper())
    p5 = cells[5].paragraphs[0]
    r5 = p5.add_run(src_label)
    r5.font.size = Pt(9)
    src_colors = {'CT': RGBColor(0xE6,0x51,0x00), 'ES': RGBColor(0x1B,0x5E,0x20),
                  'AQA': RGBColor(0x1A,0x23,0x7E)}
    r5.font.color.rgb = src_colors.get(src_label, C_DARK_BLUE)
    p5.alignment = WD_ALIGN_PARAGRAPH.CENTER

def add_question_table_header(doc, table):
    """Add header row to the question table"""
    hdr_row = table.rows[0]
    headers = ['Q.No', 'CO', 'BTL', 'Question', 'Marks', 'Source']
    widths  = [600, 600, 600, 6300, 600, 700]
    for cell, txt, w in zip(hdr_row.cells, headers, widths):
        cell.text = txt
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.runs[0] if p.runs else p.add_run(txt)
        run.text = txt
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = C_WHITE
        set_cell_bg(cell, '1A237E')
        set_cell_margins(cell)

    # Set column widths
    for i, w in enumerate(widths):
        set_col_width(table, i, w)

def make_question_table(doc):
    """Create a fresh 6-column question table"""
    table = doc.add_table(rows=1, cols=6)
    table.style = 'Table Grid'
    add_question_table_header(doc, table)
    set_table_borders(table, 'C5CAE9')
    return table

# ─── Summary stats ────────────────────────────────────────
def add_summary_page(doc, all_questions, subject_info, sources):
    add_section_heading(doc, 'Question Bank Summary', level=1)

    # Stats table
    total  = len(all_questions)
    by_type = {'SA': 0, 'MCQ': 0, 'PARTB': 0}
    by_src  = {'ct': 0, 'es': 0, 'aqa': 0}
    by_unit = {str(u): 0 for u in range(1,6)}
    by_co   = {f'CO{i}': 0 for i in range(1,6)}

    for q in all_questions:
        by_type[q.type] = by_type.get(q.type, 0) + 1
        by_src[q.source] = by_src.get(q.source, 0) + 1
        by_unit[str(q.unit)] = by_unit.get(str(q.unit), 0) + 1
        by_co[q.co]   = by_co.get(q.co, 0) + 1

    # Summary cards as a 3-col table
    table = doc.add_table(rows=2, cols=3)
    table.style = 'Table Grid'
    set_table_borders(table, 'C5CAE9')

    stats = [
        ('Total Questions', str(total), '1A237E'),
        ('Part A – Short Answer', str(by_type['SA']), '006D77'),
        ('Part A – MCQ', str(by_type['MCQ']), 'E96418'),
        ('Part B Questions', str(by_type['PARTB']), '28354F'),
        (f'Cycle Test ({", ".join(sources.get("ct",[]))})', str(by_src['ct']), '4A1486'),
        (f'End Semester ({", ".join(sources.get("es",[]))})', str(by_src['es']), '1B5E20'),
    ]
    for ri, row in enumerate(table.rows):
        for ci, cell in enumerate(row.cells):
            idx = ri * 3 + ci
            if idx >= len(stats): continue
            label, value, color = stats[idx]
            set_cell_bg(cell, 'F8F9FF')
            set_cell_margins(cell, 120, 120, 150, 150)
            p_val = cell.paragraphs[0]
            r_v = p_val.add_run(value)
            r_v.font.size = Pt(22); r_v.bold = True
            r_v.font.color.rgb = RGBColor(int(color[:2],16), int(color[2:4],16), int(color[4:],16))
            p_lbl = cell.add_paragraph()
            r_l = p_lbl.add_run(label)
            r_l.font.size = Pt(8.5)
            r_l.font.color.rgb = RGBColor(0x55,0x55,0x55)

    doc.add_paragraph()

    # Unit-wise breakdown
    add_section_heading(doc, 'Unit-wise Breakdown', level=2)
    unit_tbl = doc.add_table(rows=1, cols=6)
    unit_tbl.style = 'Table Grid'
    set_table_borders(unit_tbl, 'C5CAE9')
    for cell, txt in zip(unit_tbl.rows[0].cells, ['Unit','Topic','SA','MCQ','Part B','Total']):
        cell.text = txt
        cell.paragraphs[0].runs[0].bold = True
        cell.paragraphs[0].runs[0].font.color.rgb = C_WHITE
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_bg(cell, '283593')
        set_cell_margins(cell)

    UNIT_TOPICS = {
        '1':'Sequential Circuit Design', '2':'Asynchronous Sequential',
        '3':'Testing', '4':'PLD Design', '5':'Digital Systems'
    }
    for u in ['1','2','3','4','5']:
        uqs = [q for q in all_questions if str(q.unit) == u]
        if not uqs: continue
        row = unit_tbl.add_row()
        vals = [f'Unit {u}', UNIT_TOPICS[u],
                str(sum(1 for q in uqs if q.type=='SA')),
                str(sum(1 for q in uqs if q.type=='MCQ')),
                str(sum(1 for q in uqs if q.type=='PARTB')),
                str(len(uqs))]
        for cell, val in zip(row.cells, vals):
            cell.text = val
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_cell_margins(cell)
        set_cell_bg(row.cells[0], 'E8EAF6')
        row.cells[0].paragraphs[0].runs[0].bold = True

    set_col_width(unit_tbl, 0, 700)
    set_col_width(unit_tbl, 1, 3500)
    set_col_width(unit_tbl, 2, 900)
    set_col_width(unit_tbl, 3, 900)
    set_col_width(unit_tbl, 4, 900)
    set_col_width(unit_tbl, 5, 900)
    doc.add_paragraph()

    # CO × BTL coverage
    add_section_heading(doc, 'CO × Bloom\'s Taxonomy Coverage', level=2)
    btl_levels = ['K1','K2','K3','K4','K5','K6']
    btl_tbl = doc.add_table(rows=1, cols=8)
    btl_tbl.style = 'Table Grid'
    set_table_borders(btl_tbl, 'C5CAE9')
    for cell, txt in zip(btl_tbl.rows[0].cells, ['CO']+btl_levels+['Total']):
        cell.text = txt
        cell.paragraphs[0].runs[0].bold = True
        cell.paragraphs[0].runs[0].font.color.rgb = C_WHITE
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_bg(cell, '1A237E')
        set_cell_margins(cell)

    btl_colors_hex = {'K1':'2E7D32','K2':'1565C0','K3':'E65100','K4':'6A1B9A','K5':'C62828','K6':'37474F'}
    for co in ['CO1','CO2','CO3','CO4','CO5']:
        coqs = [q for q in all_questions if q.co == co]
        if not coqs: continue
        row = btl_tbl.add_row()
        row.cells[0].text = co
        row.cells[0].paragraphs[0].runs[0].bold = True
        set_cell_bg(row.cells[0], 'E8EAF6')
        set_cell_margins(row.cells[0])
        row.cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        for i, k in enumerate(btl_levels, 1):
            cnt = sum(1 for q in coqs if q.btl == k)
            row.cells[i].text = str(cnt) if cnt else '-'
            row.cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_cell_margins(row.cells[i])
            if cnt > 0:
                hx = btl_colors_hex.get(k,'1A237E')
                row.cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(
                    int(hx[:2],16), int(hx[2:4],16), int(hx[4:],16))
                row.cells[i].paragraphs[0].runs[0].bold = True
        row.cells[7].text = str(len(coqs))
        row.cells[7].paragraphs[0].runs[0].bold = True
        row.cells[7].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_margins(row.cells[7])

    doc.add_paragraph()

# ─── Main builder ─────────────────────────────────────────
def build_question_bank(all_questions, subject_info, sources, output_path):
    """
    all_questions : list[Question]
    subject_info  : dict with subject/code/dept/semester
    sources       : dict {'ct': ['file1.docx'], 'es': [...], 'aqa': [...]}
    output_path   : str
    """
    doc = create_document()

    generated_from = ', '.join(
        f for cat in ['ct','es','aqa'] for f in sources.get(cat,[]))

    add_document_header(doc,
        subject_info.get('subject','Digital VLSI Design'),
        subject_info.get('code','VEC311'),
        subject_info.get('dept','ECE'),
        subject_info.get('semester',''),
        generated_from)

    # CO statements (standard for VEC311)
    co_statements = {
        'CO1': 'Construct the synchronous sequential circuits.',
        'CO2': 'Solve hazards and design asynchronous sequential circuits.',
        'CO3': 'Relate the testing procedure for combinational circuit and PLA.',
        'CO4': 'Make use of PLD to construct the synchronous circuit design.',
        'CO5': 'Design and use programming tools for implementing digital circuits.',
    }
    # Override with any detected CO descriptions
    add_co_table(doc, co_statements)

    # Summary page
    add_summary_page(doc, all_questions, subject_info, sources)

    # ── Section 1: Part A – Short Answer ──────────────────
    doc.add_page_break()
    add_section_heading(doc, 'Part A – Short Answer Questions (2 Marks Each)', level=1)

    sa_qs = [q for q in all_questions if q.type == 'SA']
    for unit_no in ['1','2','3','4','5']:
        unit_qs = [q for q in sa_qs if str(q.unit) == unit_no]
        if not unit_qs: continue
        add_section_heading(doc, UNIT_NAMES.get(unit_no, f'Unit {unit_no}'), level=2)
        table = make_question_table(doc)
        for idx, q in enumerate(unit_qs, 1):
            add_question_row(table, q, idx, show_options=False)
        doc.add_paragraph()

    # ── Section 2: Part A – MCQ ───────────────────────────
    doc.add_page_break()
    add_section_heading(doc, 'Part A – MCQ Questions (2 Marks Each)', level=1)

    mcq_qs = [q for q in all_questions if q.type == 'MCQ']
    for unit_no in ['1','2','3','4','5']:
        unit_qs = [q for q in mcq_qs if str(q.unit) == unit_no]
        if not unit_qs: continue
        add_section_heading(doc, UNIT_NAMES.get(unit_no, f'Unit {unit_no}'), level=2)
        table = make_question_table(doc)
        for idx, q in enumerate(unit_qs, 1):
            add_question_row(table, q, idx, show_options=True)
        doc.add_paragraph()

    # ── Section 3: Part B ─────────────────────────────────
    doc.add_page_break()
    add_section_heading(doc, 'Part B Questions (8 / 16 Marks)', level=1)

    pb_qs = [q for q in all_questions if q.type == 'PARTB']
    for unit_no in ['1','2','3','4','5']:
        unit_qs = [q for q in pb_qs if str(q.unit) == unit_no]
        if not unit_qs: continue
        add_section_heading(doc, UNIT_NAMES.get(unit_no, f'Unit {unit_no}'), level=2)
        table = make_question_table(doc)
        for idx, q in enumerate(unit_qs, 1):
            add_question_row(table, q, idx, show_options=False)
        doc.add_paragraph()

    doc.save(output_path)
    print(f'[DOCX] Saved: {output_path}')
    return output_path


if __name__ == '__main__':
    # Quick test with parsed questions
    import sys
    sys.path.insert(0, '/home/claude')
    from qb_parser import parse_docx, Question

    files = [
        ('/mnt/user-data/uploads/VEC311_AQA_.docx', 'aqa'),
        ('/mnt/user-data/uploads/VEC311_End_Semester_Question.docx', 'es'),
        ('/mnt/user-data/uploads/VEC311-Digital_VLSI_Design_and_Technology-SET1.docx', 'ct'),
    ]

    all_qs = []
    subject_info = {}
    sources = {'ct':[],'es':[],'aqa':[]}

    for path, src in files:
        import os
        qs, info = parse_docx(path, src)
        all_qs.extend(qs)
        if not subject_info: subject_info = info
        sources[src].append(os.path.basename(path))
        print(f'{src}: {len(qs)} questions')

    # De-dup across files
    seen = set()
    unique = []
    for q in all_qs:
        key = (q.co, q.text[:50].strip().lower())
        if key not in seen:
            seen.add(key); unique.append(q)

    print(f'\nTotal unique: {len(unique)}')
    by_type = {}
    for q in unique:
        by_type[q.type] = by_type.get(q.type,0)+1
    print(f'By type: {by_type}')

    out = '/mnt/user-data/outputs/VEC311_QuestionBank.docx'
    build_question_bank(unique, subject_info, sources, out)
